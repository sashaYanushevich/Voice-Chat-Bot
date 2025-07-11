import asyncio
import json
import base64
import os
import time
import io
from typing import Dict, List
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, UploadFile, File, Form
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
import uvicorn
from dotenv import load_dotenv
import PyPDF2
from docx import Document
from fastapi.middleware.cors import CORSMiddleware

# Импорты наших компонентов
from speech_to_text import DeepgramSTT
from llm import OpenRouterClient
from aws_tts import AWSPollyTTS

load_dotenv()

app = FastAPI(title="Voice Bot API", description="Real-time voice chat bot")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# Статические файлы
app.mount("/static", StaticFiles(directory="static"), name="static")

class VoiceBotWebSocket:
    """WebSocket менеджер для голосового бота"""
    
    def __init__(self):
        self.active_connections: Dict[str, WebSocket] = {}
        self.user_sessions: Dict[str, dict] = {}
        self.cv_sessions: Dict[str, dict] = {}  # Хранилище для CV сессий
        self.user_llm_clients: Dict[str, OpenRouterClient] = {}  # Отдельный LLM клиент для каждого пользователя
        self.user_timeout_tasks: Dict[str, asyncio.Task] = {}  # Задачи таймаута для каждого пользователя
        self.user_timeout_active: Dict[str, bool] = {}  # Флаг активного таймаута для каждого пользователя
        self.user_timeout_stage: Dict[str, int] = {}  # Стадия таймаута (0, 1, 2)
        self.timeout_responses = {
            1: [  # Первая напоминалка (5 сек)
                "I didn't catch your response — would you like me to repeat the question?",
                "Are you still there?",
                "I'm waiting for your response. Are you ready to continue?"
            ],
            2: [  # Вторая напоминалка (10 сек)
                "Should I continue with the next question?",
                "Let me know when you're ready to proceed.",
                "Are you having any technical difficulties?"
            ],
            3: [  # Завершение звонка (10 сек)
                "I think we're having connection issues. Thank you for your time, and we'll be in touch soon.",
                "It seems we've lost connection. We'll contact you to reschedule the interview.",
                "Thank you for your interest. We'll reach out to you regarding next steps."
            ]
        }
        
        # Инициализируем компоненты
        try:
            print("🔧 Initializing STT...")
            self.stt = DeepgramSTT()
            print("✅ STT initialized")
            
            print("🔧 Initializing TTS...")
            self.tts = AWSPollyTTS(voice_id="Ruth", engine="generative", chunk_size=200)
            print("✅ TTS initialized")
            
            # Системный промпт
            print("🔧 Loading system prompt...")
            self.system_prompt = self._load_system_prompt()
            print("✅ System prompt loaded")
            
        except Exception as e:
            print(f"❌ Initialization error: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def _load_system_prompt(self) -> str:
        """Загружает системный промпт"""
        try:
            with open("Bot_prompt.txt", 'r', encoding='utf-8') as f:
                return f.read().strip()
        except FileNotFoundError:
            return "Ты дружелюбный AI-ассистент. Отвечай кратко и по делу на голосовые сообщения."
    
    def _extract_pdf_text(self, pdf_data: bytes) -> str:
        """Извлекает текст из PDF файла"""
        try:
            # Создаем объект BytesIO для работы с PDF
            pdf_stream = io.BytesIO(pdf_data)
            
            # Читаем PDF
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            # Извлекаем текст со всех страниц
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"
            
            return text_content.strip()
            
        except Exception as e:
            print(f"❌ PDF extraction error: {e}")
            return None
    
    def _extract_docx_text(self, docx_data: bytes) -> str:
        """Извлекает текст из DOCX файла"""
        try:
            # Создаем объект BytesIO для работы с DOCX
            docx_stream = io.BytesIO(docx_data)
            
            # Читаем DOCX
            doc = Document(docx_stream)
            
            # Извлекаем текст из всех параграфов
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content.strip()
            
        except Exception as e:
            print(f"❌ DOCX extraction error: {e}")
            return None
    
    def cleanup_old_cv_sessions(self):
        """Очищает старые CV сессии (старше 1 часа)"""
        current_time = time.time()
        old_sessions = []
        
        for session_id, session_data in self.cv_sessions.items():
            if current_time - session_data.get("uploaded_at", 0) > 3600:  # 1 час
                old_sessions.append(session_id)
        
        for session_id in old_sessions:
            del self.cv_sessions[session_id]
            print(f"🧹 Cleaned up old CV session: {session_id}")
        
        # Также очищаем неактивные LLM клиенты
        inactive_users = []
        for user_id in self.user_llm_clients.keys():
            if user_id not in self.active_connections:
                inactive_users.append(user_id)
        
        for user_id in inactive_users:
            self.user_llm_clients[user_id].clear_history()
            del self.user_llm_clients[user_id]
            print(f"🧹 Cleaned up inactive LLM client for user: {user_id}")
        
        return len(old_sessions)
    
    def get_user_llm_client(self, user_id: str) -> OpenRouterClient:
        """Получает или создает LLM клиент для конкретного пользователя"""
        if user_id not in self.user_llm_clients:
            print(f"🔧 Creating new LLM client for user {user_id}")
            self.user_llm_clients[user_id] = OpenRouterClient()
        return self.user_llm_clients[user_id]
    
    async def start_response_timeout(self, user_id: str):
        """Запускает таймер ожидания ответа пользователя"""
        # Проверяем, не активен ли уже таймер
        if self.user_timeout_active.get(user_id, False):
            print(f"⏰ Timeout already active for {user_id}, skipping")
            return
            
        # Отменяем предыдущий таймер если есть
        if user_id in self.user_timeout_tasks:
            self.user_timeout_tasks[user_id].cancel()
        
        # Устанавливаем флаг активного таймера и начальную стадию
        self.user_timeout_active[user_id] = True
        self.user_timeout_stage[user_id] = 1  # Начинаем с первой стадии
        
        # Создаем новый таймер (первая стадия - 5 секунд)
        self.user_timeout_tasks[user_id] = asyncio.create_task(
            self._timeout_handler(user_id)
        )
    
    async def cancel_response_timeout(self, user_id: str):
        """Отменяет таймер ожидания ответа"""
        # Сбрасываем флаг активного таймера и стадию
        self.user_timeout_active[user_id] = False
        self.user_timeout_stage[user_id] = 0
        
        if user_id in self.user_timeout_tasks:
            self.user_timeout_tasks[user_id].cancel()
            del self.user_timeout_tasks[user_id]
    
    async def _timeout_handler(self, user_id: str):
        """Обработчик таймаута с тремя стадиями"""
        try:
            stage = self.user_timeout_stage.get(user_id, 1)
            
            # Определяем время ожидания в зависимости от стадии
            if stage == 1:
                wait_time = 15.0  # Первая напоминалка через 5 сек
            elif stage == 2:
                wait_time = 25.0  # Вторая напоминалка через 10 сек
            else:  # stage == 3
                wait_time = 30.0  # Завершение через 10 сек
            
            print(f"⏰ Starting timeout stage {stage} for {user_id} (waiting {wait_time}s)")
            await asyncio.sleep(wait_time)
            
            # Проверяем, что таймер все еще активен
            if not self.user_timeout_active.get(user_id, False):
                return
            
            # Если дошли сюда, значит пользователь не ответил
            import random
            timeout_message = random.choice(self.timeout_responses[stage])
            
            print(f"⏰ Timeout stage {stage} for user {user_id}: {timeout_message}")
            
            await self.send_message(user_id, {
                "type": "bot_text",
                "text": timeout_message
            })
            
            # Озвучиваем сообщение о таймауте
            if stage == 3:
                await self.send_message(user_id, {
                    "type": "status",
                    "message": "🔊 HR is ending the interview..."
                })
            else:
                await self.send_message(user_id, {
                    "type": "status",
                    "message": "🔊 HR is prompting you..."
                })
            
            await self._synthesize_and_send_audio(user_id, timeout_message)
            
            if stage == 3:
                # Завершаем интервью
                await self.send_message(user_id, {
                    "type": "interview_ended",
                    "message": "🔚 Interview ended due to no response"
                })
                
                # Отключаем пользователя
                print(f"🔚 Ending interview for {user_id} due to no response")
                self.disconnect(user_id)
                return
            else:
                await self.send_message(user_id, {
                    "type": "completed",
                    "message": "✅ Ready for your response"
                })
            
            # Переходим к следующей стадии
            self.user_timeout_stage[user_id] = stage + 1
            
            # ВАЖНО: НЕ сбрасываем флаг здесь - оставляем активным до audio_playback_complete
            # Это предотвратит запуск новых таймеров пока играет timeout сообщение
            
        except asyncio.CancelledError:
            # Таймер был отменен (пользователь ответил вовремя)
            self.user_timeout_active[user_id] = False
            self.user_timeout_stage[user_id] = 0
        except Exception as e:
            print(f"❌ Timeout handler error for {user_id}: {e}")
            self.user_timeout_active[user_id] = False
            self.user_timeout_stage[user_id] = 0
    
    async def connect(self, websocket: WebSocket, user_id: str, session_id: str = None):
        """Подключение нового пользователя"""
        await websocket.accept()
        self.active_connections[user_id] = websocket
        self.user_sessions[user_id] = {
            "conversation_history": [],
            "connected_at": time.time(),
            "cv_text": None,
            "candidate_info": None
        }
        
        # Очищаем старые сессии
        self.cleanup_old_cv_sessions()
        
        # Загружаем данные CV если есть session_id
        if session_id and session_id in self.cv_sessions:
            session_data = self.cv_sessions[session_id]
            self.user_sessions[user_id]["cv_text"] = session_data["cv_text"]
            self.user_sessions[user_id]["candidate_info"] = session_data["candidate_info"]
            
            candidate = session_data["candidate_info"]
            print(f"📄 Loaded CV data for {candidate['firstName']} {candidate['lastName']}")
        
        await self.send_message(user_id, {
            "type": "connected",
            "message": "🎤 HR Interview starting!"
        })
        
        print(f"✅ User {user_id} connected")
        
        # Автоматически начинаем интервью
        await self.start_interview(user_id)
    
    async def start_interview(self, user_id: str):
        """Автоматически начинает интервью с приветствием HR"""
        try:
            print(f"🎬 Starting automatic interview for {user_id}")
            
            # Генерируем приветственное сообщение от HR
            greeting_prompt = "Start the interview exactly as instructed in the prompt. Follow the 'Begin with:' instruction precisely."
            
            # Создаем расширенный системный промпт с CV и данными кандидата если есть
            enhanced_prompt = self.system_prompt
            if user_id in self.user_sessions:
                session = self.user_sessions[user_id]
                if session.get("cv_text") or session.get("candidate_info"):
                    enhanced_prompt = f"""{self.system_prompt}

CANDIDATE INFORMATION:"""
                    
                    if session.get("candidate_info"):
                        candidate = session["candidate_info"]
                        enhanced_prompt += f"""
Name: {candidate.get('firstName', '')} {candidate.get('lastName', '')}
Email: {candidate.get('email', '')}"""
                    
                    if session.get("cv_text"):
                        enhanced_prompt += f"""

CV CONTENT:
{session['cv_text']}"""
                    
                    enhanced_prompt += """

Use this information to conduct a personalized interview, asking relevant questions based on their CV and experience.

CRITICAL: Keep response under 30 words. Be extremely brief and direct."""
            
            # Получаем персональный LLM клиент для пользователя
            user_llm = self.get_user_llm_client(user_id)
            bot_response = await user_llm.chat_completion(greeting_prompt, enhanced_prompt)
            print(f"🤖 HR greeting to {user_id}: {bot_response}")
            
            await self.send_message(user_id, {
                "type": "bot_text", 
                "text": bot_response
            })
            
            # Озвучиваем приветствие
            await self.send_message(user_id, {
                "type": "status",
                "message": "🔊 HR is greeting you..."
            })
            
            await self._synthesize_and_send_audio(user_id, bot_response)
            
            await self.send_message(user_id, {
                "type": "completed",
                "message": "✅ Ready to hear your response"
            })
            
            # НЕ запускаем таймер здесь - ждем уведомления о завершении воспроизведения
            # await self.start_response_timeout(user_id)
            
        except Exception as e:
            print(f"❌ Error starting interview for {user_id}: {e}")
            await self.send_message(user_id, {
                "type": "error",
                "message": "❌ Error starting interview"
            })
    
    def disconnect(self, user_id: str):
        """Отключение пользователя"""
        if user_id in self.active_connections:
            del self.active_connections[user_id]
        if user_id in self.user_sessions:
            del self.user_sessions[user_id]
        if user_id in self.user_llm_clients:
            # Очищаем историю разговора перед удалением
            self.user_llm_clients[user_id].clear_history()
            del self.user_llm_clients[user_id]
            print(f"🧹 Cleared LLM client for user {user_id}")
        # Отменяем таймер ожидания если есть
        if user_id in self.user_timeout_tasks:
            self.user_timeout_tasks[user_id].cancel()
            del self.user_timeout_tasks[user_id]
            print(f"🧹 Cancelled timeout task for user {user_id}")
        # Очищаем флаги таймаута
        if user_id in self.user_timeout_active:
            del self.user_timeout_active[user_id]
        if user_id in self.user_timeout_stage:
            del self.user_timeout_stage[user_id]
        print(f"❌ User {user_id} disconnected")
    
    async def send_message(self, user_id: str, message: dict):
        """Отправка сообщения пользователю"""
        if user_id in self.active_connections:
            try:
                await self.active_connections[user_id].send_text(json.dumps(message, ensure_ascii=False))
            except Exception as e:
                print(f"❌ Send error to {user_id}: {e}")
                self.disconnect(user_id)
    
    async def process_audio(self, user_id: str, audio_data: bytes):
        """Обработка аудио сообщения"""
        print(f"🎬 Starting audio processing for {user_id}")
        try:
            # 1. STT - распознаем речь
            print(f"🎧 Starting STT for {user_id}")
            await self.send_message(user_id, {
                "type": "status",
                "message": "🎧 Recognizing speech..."
            })
            
            user_text = await self.stt.transcribe_audio_bytes(audio_data)
            print(f"🎧 STT result for {user_id}: '{user_text}'")
            
            if not user_text.strip():
                print(f"⚠️ Empty STT result for {user_id} - ignoring, keeping timeout active")
                await self.send_message(user_id, {
                    "type": "status",
                    "message": "🎧 Could not recognize speech, please try again"
                })
                # НЕ отменяем таймер для пустых результатов - продолжаем ждать
                return
            
            # Отменяем таймер ожидания только если получили реальный текст
            await self.cancel_response_timeout(user_id)
            
            print(f"👤 User {user_id}: {user_text}")
            
            await self.send_message(user_id, {
                "type": "user_text",
                "text": user_text
            })
            
            # 2. LLM - генерируем ответ
            print(f"🧠 Starting LLM for {user_id}")
            await self.send_message(user_id, {
                "type": "status",
                "message": "🧠 Thinking about response..."
            })
            
            # Создаем расширенный системный промпт с CV и данными кандидата если есть
            enhanced_prompt = self.system_prompt
            if user_id in self.user_sessions:
                session = self.user_sessions[user_id]
                if session.get("cv_text") or session.get("candidate_info"):
                    enhanced_prompt = f"""{self.system_prompt}

CANDIDATE INFORMATION:"""
                    
                    if session.get("candidate_info"):
                        candidate = session["candidate_info"]
                        enhanced_prompt += f"""
Name: {candidate.get('firstName', '')} {candidate.get('lastName', '')}
Email: {candidate.get('email', '')}"""
                    
                    if session.get("cv_text"):
                        enhanced_prompt += f"""

CV CONTENT:
{session['cv_text']}"""
                    
                    enhanced_prompt += """

Use this information to conduct a personalized interview, asking relevant questions based on their CV and experience.

CRITICAL: Keep response under 30 words. Be extremely brief and direct."""
                
                # Получаем персональный LLM клиент для пользователя
                user_llm = self.get_user_llm_client(user_id)
                bot_response = await user_llm.chat_completion(user_text, enhanced_prompt)
            
            print(f"🧠 LLM result for {user_id}: '{bot_response}'")
            
            print(f"🤖 Bot to {user_id}: {bot_response}")
            
            await self.send_message(user_id, {
                "type": "bot_text",
                "text": bot_response
            })
            
            # 3. TTS - озвучиваем ответ
            print(f"🔊 Starting TTS for {user_id}")
            await self.send_message(user_id, {
                "type": "status",
                "message": "🔊 Generating speech..."
            })
            
            # Генерируем аудио чанками и отправляем
            await self._synthesize_and_send_audio(user_id, bot_response)
            
            await self.send_message(user_id, {
                "type": "completed",
                "message": "✅ Ready for next question"
            })
            
            # НЕ запускаем таймер здесь - ждем уведомления о завершении воспроизведения
            # await self.start_response_timeout(user_id)
            
            print(f"✅ Audio processing completed for {user_id}")
            
        except Exception as e:
            print(f"❌ Process error for {user_id}: {e}")
            import traceback
            traceback.print_exc()
            await self.send_message(user_id, {
                "type": "error",
                "message": f"❌ Processing error: {str(e)}"
            })
    
    async def _synthesize_and_send_audio(self, user_id: str, text: str):
        """Синтез речи с отправкой чанков"""
        try:
            # Разбиваем на чанки
            chunks = self.tts.split_text_into_chunks(text)
            
            for i, chunk in enumerate(chunks):
                # Синтезируем чанк
                audio_data = await self.tts.synthesize_chunk(chunk)
                if audio_data:
                    # Кодируем в base64 и отправляем
                    audio_b64 = base64.b64encode(audio_data).decode('utf-8')
                    await self.send_message(user_id, {
                        "type": "audio_chunk",
                        "audio": audio_b64,
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    })
                    
                    # Небольшая задержка между чанками
                    await asyncio.sleep(0.1)
                    
        except Exception as e:
            print(f"❌ TTS error for {user_id}: {e}")
            await self.send_message(user_id, {
                "type": "error",
                "message": "❌ Speech synthesis error"
            })
    
    async def process_cv_upload(self, user_id: str, filename: str, base64_data: str, candidate_info: dict = None):
        """Обработка загрузки CV в PDF"""
        print(f"📄 Processing CV upload for {user_id}: {filename}")
        
        try:
            # Декодируем base64 данные
            pdf_data = base64.b64decode(base64_data)
            print(f"📄 PDF size: {len(pdf_data)} bytes")
            
            # Извлекаем текст из PDF
            cv_text = self._extract_pdf_text(pdf_data)
            
            if not cv_text:
                await self.send_message(user_id, {
                    "type": "cv_error",
                    "message": "Could not extract text from PDF"
                })
                return
            
            # Сохраняем CV и данные кандидата в сессии пользователя
            if user_id in self.user_sessions:
                self.user_sessions[user_id]["cv_text"] = cv_text
                self.user_sessions[user_id]["candidate_info"] = candidate_info
            
            print(f"📄 CV extracted for {user_id}: {len(cv_text)} characters")
            print(f"📄 CV preview: {cv_text[:200]}...")
            
            if candidate_info:
                print(f"👤 Candidate: {candidate_info.get('firstName')} {candidate_info.get('lastName')} ({candidate_info.get('email')})")
            
            await self.send_message(user_id, {
                "type": "cv_uploaded",
                "message": f"CV uploaded successfully ({len(cv_text)} characters)"
            })
            
        except Exception as e:
            print(f"❌ CV processing failed for {user_id}: {e}")
            import traceback
            traceback.print_exc()
            await self.send_message(user_id, {
                "type": "cv_error",
                "message": "Error processing CV"
            })

# Создаем глобальный экземпляр
voice_bot = VoiceBotWebSocket()

@app.websocket("/ws/{user_id}")
async def websocket_endpoint(websocket: WebSocket, user_id: str):
    """WebSocket эндпоинт для голосового чата"""
    # Получаем session_id из query параметров
    query_params = dict(websocket.query_params)
    session_id = query_params.get('session_id')
    
    await voice_bot.connect(websocket, user_id, session_id)
    
    try:
        while True:
            # Получаем сообщение от клиента
            message = await websocket.receive_text()
            data = json.loads(message)
            
            # Проверяем, что пользователь все еще подключен
            if user_id not in voice_bot.active_connections:
                print(f"⚠️ Ignoring message from disconnected user {user_id}")
                break
            
            print(f"📨 Received message from {user_id}: type={data.get('type')}")
            
            if data["type"] == "audio":
                print(f"🎵 Processing audio from {user_id}, size: {len(data['audio'])} chars")
                # Декодируем аудио из base64
                try:
                    audio_data = base64.b64decode(data["audio"])
                    print(f"🔓 Decoded audio: {len(audio_data)} bytes")
                    await voice_bot.process_audio(user_id, audio_data)
                except Exception as e:
                    print(f"❌ Audio decode error: {e}")
                    await voice_bot.send_message(user_id, {
                        "type": "error",
                        "message": f"❌ Ошибка декодирования аудио: {str(e)}"
                    })
                

                
            elif data["type"] == "ping":
                await voice_bot.send_message(user_id, {"type": "pong"})
                
            elif data["type"] == "audio_playback_complete":
                print(f"🔊 Audio playback completed for {user_id}")
                
                # Если таймер был активен (timeout сообщение закончилось)
                if voice_bot.user_timeout_active.get(user_id, False):
                    current_stage = voice_bot.user_timeout_stage.get(user_id, 1)
                    print(f"🔊 Timeout message playback completed for {user_id}, stage {current_stage-1}")
                    
                    # Если это была третья стадия (завершение интервью), не запускаем новый таймер
                    if current_stage > 3:
                        print(f"🔚 Interview ended for {user_id}, not starting new timeout")
                        return
                    
                    # Сбрасываем флаг активного таймера
                    voice_bot.user_timeout_active[user_id] = False
                    
                    # Запускаем следующую стадию таймера
                    print(f"🔊 Starting timeout stage {current_stage} for {user_id}")
                    voice_bot.user_timeout_tasks[user_id] = asyncio.create_task(
                        voice_bot._timeout_handler(user_id)
                    )
                    voice_bot.user_timeout_active[user_id] = True
                else:
                    # Обычное завершение воспроизведения (не timeout сообщение)
                    print(f"🔊 Starting new timeout for {user_id}")
                    await voice_bot.start_response_timeout(user_id)
                
    except WebSocketDisconnect:
        voice_bot.disconnect(user_id)
    except Exception as e:
        print(f"❌ WebSocket error: {e}")
        voice_bot.disconnect(user_id)

@app.post("/upload-cv")
async def upload_cv(
    cv_file: UploadFile = File(...),
    first_name: str = Form(...),
    last_name: str = Form(...),
    email: str = Form(...)
):
    """Загрузка CV и данных кандидата"""
    try:
        # Проверяем тип файла
        allowed_types = [
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        if cv_file.content_type not in allowed_types:
            return JSONResponse(
                status_code=400,
                content={"error": "Only PDF and DOCX files are allowed"}
            )
        
        # Проверяем размер файла (10MB)
        contents = await cv_file.read()
        if len(contents) > 10 * 1024 * 1024:
            return JSONResponse(
                status_code=400,
                content={"error": "File size must be less than 10MB"}
            )
        
        # Извлекаем текст в зависимости от типа файла
        if cv_file.content_type == "application/pdf":
            cv_text = voice_bot._extract_pdf_text(contents)
            file_type = "PDF"
        elif cv_file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            cv_text = voice_bot._extract_docx_text(contents)
            file_type = "DOCX"
        else:
            return JSONResponse(
                status_code=400,
                content={"error": "Unsupported file type"}
            )
        
        if not cv_text:
            return JSONResponse(
                status_code=400,
                content={"error": f"Could not extract text from {file_type} file"}
            )
        
        # Генерируем уникальный ID для сессии
        session_id = f"session_{int(time.time())}_{hash(email)}"
        
        # Сохраняем данные в временном хранилище
        voice_bot.cv_sessions[session_id] = {
            "candidate_info": {
                "firstName": first_name,
                "lastName": last_name,
                "email": email
            },
            "cv_text": cv_text,
            "uploaded_at": time.time()
        }
        
        print(f"📄 CV uploaded for {first_name} {last_name} ({email})")
        print(f"📄 Session ID: {session_id}")
        print(f"📄 CV length: {len(cv_text)} characters")
        
        return JSONResponse(content={
            "success": True,
            "session_id": session_id,
            "message": f"CV uploaded successfully ({len(cv_text)} characters)",
            "candidate": {
                "firstName": first_name,
                "lastName": last_name,
                "email": email
            }
        })
        
    except Exception as e:
        print(f"❌ CV upload error: {e}")
        import traceback
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"error": "Server error processing CV"}
        )

@app.get("/")
async def get_index():
    """Главная страница"""
    return FileResponse("static/index.html")

@app.get("/health")
async def health_check():
    """Проверка здоровья сервиса"""
    return {"status": "OK", "message": "Voice Bot API is running"}

@app.get("/test-components")
async def test_components():
    """Тест всех компонентов API"""
    results = {}
    
    try:
        # Тест STT (с заглушкой)
        results["stt"] = "✅ STT component initialized"
        
        # Тест LLM
        test_llm = OpenRouterClient()
        test_response = await test_llm.chat_completion("Say hello", voice_bot.system_prompt)
        results["llm"] = f"✅ LLM response: {test_response[:50]}..."
        
        # Тест TTS
        chunks = voice_bot.tts.split_text_into_chunks("Hello world")
        results["tts"] = f"✅ TTS chunks: {len(chunks)}"
        
    except Exception as e:
        results["error"] = f"❌ Component test failed: {str(e)}"
    
    return results

if __name__ == "__main__":
    # Создаем папку static если её нет
    os.makedirs("static", exist_ok=True)
    
    
    uvicorn.run(
        "websocket_api:app",
        host="0.0.0.0",
        port=8800,
        reload=True,
    ) 